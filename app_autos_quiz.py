import os
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone

import pandas as pd
import streamlit as st
from filelock import FileLock, Timeout

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

DATA_FILE = Path("autos_submissions.csv")
LOCK_FILE = Path("autos_submissions.csv.lock")

st.set_page_config(page_title="DATA-101 Automobile Quiz (Pandas) — Team Game", layout="wide")

def init_storage():
    if not DATA_FILE.exists():
        cols = ["ts_iso","round","team","q1_rows","q1_cols","q1_score",
                "q2_company","q2_company_score","q2_price","q2_price_score",
                "q3_last_price","q3_score",
                "q4_bodystyle","q4_price","q4_score",
                "total_score"]
        pd.DataFrame(columns=cols).to_csv(DATA_FILE, index=False)

def write_submission(row: dict):
    init_storage()
    lock = FileLock(str(LOCK_FILE))
    try:
        with lock.acquire(timeout=5):
            df = pd.read_csv(DATA_FILE)
            df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
            df.to_csv(DATA_FILE, index=False)
    except Timeout:
        st.error("Server is busy. Please try submitting again.")

def read_submissions() -> pd.DataFrame:
    init_storage()
    return pd.read_csv(DATA_FILE)

@st.cache_data
def load_csv(file) -> pd.DataFrame:
    return pd.read_csv(file)

def compute_answers(df: pd.DataFrame, company_col: str, price_col: str, body_col: str):
    # 1) rows, cols
    nrows, ncols = df.shape

    # coerce price to numeric
    prices = pd.to_numeric(df[price_col], errors="coerce")

    # 2) most expensive car company + price
    idxmax = prices.idxmax()
    company_of_max = None if pd.isna(idxmax) else str(df.loc[idxmax, company_col])
    max_price_value = None if pd.isna(idxmax) else float(prices.loc[idxmax])

    # 3) last observation price
    last_price = float(pd.to_numeric(df.loc[df.index.max(), price_col], errors="coerce"))

    # 4) body-style with highest max price + value
    by_body_max = df.groupby(body_col)[price_col].apply(lambda s: pd.to_numeric(s, errors="coerce").max())
    top_body = str(by_body_max.idxmax())
    top_body_price = float(by_body_max.max())

    return {
        "shape": (int(nrows), int(ncols)),
        "company_of_max": company_of_max,
        "max_price_value": max_price_value,
        "last_price": last_price,
        "top_body": top_body,
        "top_body_price": top_body_price,
        "by_body_max_table": by_body_max.sort_values(ascending=False).reset_index().rename(columns={price_col: "max_price"})
    }

def make_docx(answers: dict, mapping: dict) -> bytes:
    doc = Document()
    doc.add_heading("DATA-101 — Automobile Quiz (Answer Key)", level=0)
    doc.add_paragraph(f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("Dataset Mapping", level=1)
    for k,v in mapping.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Answers", level=1)
    r,c = answers["shape"]
    doc.add_paragraph(f"1) Rows, Columns: {r}, {c}")
    doc.add_paragraph(f"2) Most expensive car (company): {answers['company_of_max']} at price {answers['max_price_value']}")
    doc.add_paragraph(f"3) Price of last observation: {answers['last_price']}")
    doc.add_paragraph(f"4) Body-style with highest max price: {answers['top_body']} at price {answers['top_body_price']}")

    doc.add_heading("Max Price by Body-Style", level=2)
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = "Body-Style"
    hdr[1].text = "Max Price"
    for _,row in answers["by_body_max_table"].iterrows():
        cells = t.add_row().cells
        cells[0].text = str(row.iloc[0])
        cells[1].text = str(row.iloc[1])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def make_pdf(answers: dict, mapping: dict) -> bytes:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "DATA-101 — Automobile Quiz (Answer Key)")
    y -= 20
    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    y -= 30
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Dataset Mapping")
    y -= 18
    c.setFont("Helvetica", 10)
    for k,v in mapping.items():
        c.drawString(60, y, f"{k}: {v}")
        y -= 14

    y -= 10
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Answers")
    y -= 18
    c.setFont("Helvetica", 10)
    r,cnt = answers["shape"]
    lines = [
        f"1) Rows, Columns: {r}, {cnt}",
        f"2) Most expensive car (company): {answers['company_of_max']} at price {answers['max_price_value']}",
        f"3) Price of last observation: {answers['last_price']}",
        f"4) Body-style with highest max price: {answers['top_body']} at price {answers['top_body_price']}"
    ]
    for line in lines:
        c.drawString(60, y, line); y -= 14
        if y < 60: c.showPage(); y = height - 50

    y -= 10
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Max Price by Body-Style")
    y -= 18
    c.setFont("Helvetica", 10)
    for _,row in answers["by_body_max_table"].iterrows():
        c.drawString(60, y, f"{row.iloc[0]}: {row.iloc[1]}"); y -= 14
        if y < 60: c.showPage(); y = height - 50

    c.showPage()
    c.save()
    bio.seek(0)
    return bio.read()

def main():
    st.title("🏎️ DATA-101 Automobile Quiz — Team Game (Pandas fundamentals)")
    st.caption("Upload the Automobile CSV, map columns, and let teams compete. The app derives answers and auto-scores.")

    # ===== Instructor setup =====
    with st.expander("🧑‍🏫 Instructor Setup", expanded=True):
        uploaded = st.file_uploader("Upload the Automobile CSV", type=["csv"])
        if uploaded is not None:
            df = load_csv(uploaded)
            st.success(f"Loaded CSV with shape {df.shape}")
            st.dataframe(df.head(), use_container_width=True)

            cols = list(df.columns)
            c1, c2, c3 = st.columns(3)
            with c1:
                company_col = st.selectbox("Company column", cols, index=0)
            with c2:
                price_col = st.selectbox("Price column", cols, index=min(1, len(cols)-1))
            with c3:
                body_col = st.selectbox("Body-style column", cols, index=min(2, len(cols)-1))

            answers = compute_answers(df, company_col, price_col, body_col)
            st.info("Derived correct answers from your dataset.")
            st.write({"shape": answers["shape"],
                      "company_of_max": answers["company_of_max"],
                      "max_price_value": answers["max_price_value"],
                      "last_price": answers["last_price"],
                      "top_body": answers["top_body"],
                      "top_body_price": answers["top_body_price"]})

            mapping = {"Company column": company_col, "Price column": price_col, "Body-style column": body_col}
            colA, colB = st.columns(2)
            with colA:
                data = make_docx(answers, mapping)
                st.download_button("⬇️ Download Answer Key (DOCX)", data=data,
                                   file_name="Automobile_Quiz_Answer_Key.docx")
            with colB:
                data = make_pdf(answers, mapping)
                st.download_button("⬇️ Download Answer Key (PDF)", data=data,
                                   file_name="Automobile_Quiz_Answer_Key.pdf")
            # store mapping in session for the game section
            st.session_state["mapping"] = mapping
            st.session_state["uploaded_bytes"] = uploaded.getvalue()
        else:
            st.warning("Upload a CSV to enable the quiz.")

    st.divider()

    # Stop if no CSV yet
    if "uploaded_bytes" not in st.session_state:
        st.info("Waiting for CSV upload...")
        st.stop()

    # ===== Game play =====
    st.subheader("📝 Student Submission")
    df = load_csv(BytesIO(st.session_state["uploaded_bytes"]))
    mapping = st.session_state["mapping"]
    company_col, price_col, body_col = mapping["Company column"], mapping["Price column"], mapping["Body-style column"]
    answers = compute_answers(df, company_col, price_col, body_col)

    left, right = st.columns([2,1], gap="large")
    with left:
        with st.form("quiz_form", clear_on_submit=True):
            team = st.text_input("Team name", placeholder="e.g., Data Warriors")
            round_num = st.number_input("Round #", 1, 10, value=1, step=1)

            st.markdown("**Q1) Using pandas, how many rows and columns are in the dataset?**")
            q1_rows = st.number_input("Rows", min_value=0, value=0, step=1)
            q1_cols = st.number_input("Columns", min_value=0, value=0, step=1)

            st.markdown("**Q2) Report the most expensive car price based on the company.**")
            companies = sorted(df[company_col].astype(str).unique().tolist())
            q2_company = st.selectbox("Which company has the single most expensive car?", companies)
            q2_price = st.number_input("What is that car's price?", min_value=0.0, step=100.0, value=0.0, format="%.2f")

            st.markdown("**Q3) Report the price of the last observation in the dataframe.**")
            q3_last = st.number_input("Last observation price", min_value=0.0, step=100.0, value=0.0, format="%.2f")

            st.markdown("**Q4) Report the most expensive car price based on body-style.**")
            bodies = sorted(df[body_col].astype(str).unique().tolist())
            q4_body = st.selectbox("Which body-style has the highest max price?", bodies)
            q4_price = st.number_input("What is that max price for the chosen body-style?", min_value=0.0, step=100.0, value=0.0, format="%.2f")

            submitted = st.form_submit_button("Submit answers 🚀")

        if submitted:
            if not team.strip():
                st.warning("Enter a team name.")
            else:
                s_q1 = int((q1_rows == answers["shape"][0]) and (q1_cols == answers["shape"][1])) * 3
                s_q2_company = 1 if q2_company == answers["company_of_max"] else 0
                s_q2_price = 2 if abs(q2_price - answers["max_price_value"]) < 1e-6 else 0
                s_q3 = 3 if abs(q3_last - answers["last_price"]) < 1e-6 else 0
                s_q4 = 1 if q4_body == answers["top_body"] else 0
                s_q4 += 2 if abs(q4_price - answers["top_body_price"]) < 1e-6 else 0

                total = s_q1 + s_q2_company + s_q2_price + s_q3 + s_q4

                row = {
                    "ts_iso": datetime.now(timezone.utc).isoformat(),
                    "round": int(round_num),
                    "team": team.strip(),
                    "q1_rows": int(q1_rows), "q1_cols": int(q1_cols), "q1_score": s_q1,
                    "q2_company": q2_company, "q2_company_score": s_q2_company,
                    "q2_price": float(q2_price), "q2_price_score": s_q2_price,
                    "q3_last_price": float(q3_last), "q3_score": s_q3,
                    "q4_bodystyle": q4_body, "q4_price": float(q4_price), "q4_score": s_q4,
                    "total_score": total
                }
                write_submission(row)
                st.success(f"Submitted! Score = {total}")

    with right:
        st.subheader("📊 Live Leaderboard")
        df_sub = read_submissions()
        if df_sub.empty:
            st.info("No submissions yet.")
        else:
            latest_round = int(df_sub["round"].max())
            show_round = st.number_input("Leaderboard round", 1, max(1, latest_round), value=latest_round, step=1)
            view = df_sub[df_sub["round"] == show_round].copy()
            view["ts_iso"] = pd.to_datetime(view["ts_iso"])
            view.sort_values(["team","total_score","ts_iso"], ascending=[True, False, True], inplace=True)
            best = view.groupby("team", as_index=False).first()
            best.sort_values(["total_score","ts_iso"], ascending=[False, True], inplace=True)
            best.insert(0, "rank", range(1, len(best)+1))
            best.rename(columns={"ts_iso":"submitted_utc"}, inplace=True)
            st.dataframe(best[["rank","team","total_score","submitted_utc"]], use_container_width=True, hide_index=True)

            with st.expander("All submissions (this round)"):
                st.dataframe(view.sort_values("ts_iso"), use_container_width=True, hide_index=True)

    st.divider()
    st.subheader("🧰 Admin")
    with st.form("admin"):
        action = st.selectbox("Action", ["(choose)", "Reset all data"])
        code = st.text_input("Admin code", type="password", placeholder="enter the code")
        go = st.form_submit_button("Apply")
    admin_code = os.getenv("ADMIN_CODE", "letmein")
    if go:
        if code != admin_code:
            st.error("Wrong admin code.")
        else:
            if action == "Reset all data":
                try:
                    lock = FileLock(str(LOCK_FILE))
                    with lock.acquire(timeout=5):
                        if DATA_FILE.exists():
                            DATA_FILE.unlink()
                        init_storage()
                    st.success("All submissions cleared.")
                except Timeout:
                    st.error("Could not acquire file lock. Try again.")
            else:
                st.info("Choose an action first.")

if __name__ == "__main__":
    init_storage()
    main()
